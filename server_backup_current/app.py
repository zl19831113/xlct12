# 初始化数据库
def init_db():
    """初始化数据库表结构"""
    db.create_all()
    print("数据库表结构已初始化")

def upgrade_db():
    """升级数据库结构和数据"""
    with app.app_context():
        # 检查是否需要添加 source_type 字段
        try:
            # 尝试查询以检查字段是否存在
            Paper.query.with_entities(Paper.source_type).first()
            print("source_type 字段已存在")
        except Exception as e:
            if 'no such column' in str(e).lower():
                # 字段不存在，需要添加
                try:
                    # 使用原生SQL添加字段
                    db.session.execute(text("ALTER TABLE papers ADD COLUMN source_type VARCHAR(20) DEFAULT '地区联考'"))
                    db.session.commit()
                    
                    # 根据来源名称智能判断来源类型并更新
                    papers = Paper.query.all()
                    for paper in papers:
                        # 如果来源中包含"中学"、"附属"、"大学"等关键词，则设为名校调考
                        if any(keyword in paper.source for keyword in ['中学', '附属', '大学', '一中', '二中', '三中', '四中', '五中', '六中', '七中', '八中', '九中', '十中']):
                            paper.source_type = '名校调考'
                        else:
                            paper.source_type = '地区联考'
                    
                    db.session.commit()
                    print("成功添加 source_type 字段并智能更新现有数据")
                except Exception as add_error:
                    db.session.rollback()
                    print(f"添加 source_type 字段失败: {str(add_error)}")

# Flask 3.x不再支持before_first_request
def check_db_tables():
    """在应用启动时检查数据库表结构"""
    try:
        with app.app_context():
            # 检查是否需要创建表
            inspector = db.inspect(db.engine)
            if not inspector.has_table('papers'):
                init_db()
                print("数据库表格已创建")
            else:
                print("数据库表格已存在")
                
            # 升级数据库结构
            upgrade_db()
    except Exception as e:
        print(f"数据库检查/创建时出错: {str(e)}")

import os
import sqlite3
import datetime
import unicodedata
import csv
import re
import random
import json
import traceback  # 添加traceback模块的导入
import logging
import hashlib
import io
import base64
import glob
import math
import uuid
import re
import io
import math
import zipfile  # 添加zipfile模块的导入
from flask import Flask, render_template, request, jsonify, send_file, redirect, url_for, abort
from flask_sqlalchemy import SQLAlchemy
from flask_wtf.csrf import CSRFProtect
from sqlalchemy import text
from werkzeug.utils import secure_filename
import base64
from docx import Document
from io import BytesIO
from docx.shared import Inches, RGBColor
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.style import WD_STYLE_TYPE
import re
from docx import Document
from docx.shared import Pt
from datetime import datetime, timedelta
from urllib.parse import unquote
from tqdm import tqdm
import numpy as np
import pandas as pd
import traceback
from math import radians, cos, sin, asin, sqrt
from functools import wraps

# 设置表格边框的辅助函数
def set_cell_border(table):
    """Set cell borders for Word table"""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    
    # 设置表格的全部单元格边框
    tbl = table._tbl
    for row in table.rows:
        for cell in row.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            # 给单元格添加边框
            tcBorders = OxmlElement('w:tcBorders')
            for border_name in ['top', 'left', 'bottom', 'right']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), '4')  # 1pt = 8 units
                border.set(qn('w:space'), '0')
                border.set(qn('w:color'), '000000')  # 黑色边框
                tcBorders.append(border)
            
            tcPr.append(tcBorders)

def optimize_choice_display(choice_text: str):
    """
    将类似：
      2024年12月3日，亭山商代遗址群出土青铜器展……见证了（ ） ①农业出现… ②人类摆脱… ③在一定生产关系… ④封建社会… A．①② B．①④ C．②③ D．③④
    变成：
      2024年12月3日，亭山商代遗址群出土青铜器展……见证了（ ）
      ①农业出现…
      ②人类摆脱…
      ③在一定生产关系…
      ④封建社会…
      A．①② B．①④ C．②③ D．③④
    """

    if not choice_text:
        return []

    # 1) 合并多余空格、换行
    choice_text = re.sub(r'[\r\n]+', ' ', choice_text)
    choice_text = re.sub(r'\s+', ' ', choice_text).strip()
    
    # 移除任何剩余的数字编号格式，如 "3．" 或类似格式
    choice_text = re.sub(r'\d+[．.、]\s*', '', choice_text)

    # 2) 在 "（ ）" 后插入换行
    choice_text = re.sub(r'（\s*）', '（ ）\n', choice_text, count=1)

    # 3) 把 ① ② ③ ④ 分别放在新行
    for bullet in ['①', '②', '③', '④']:
        choice_text = choice_text.replace(bullet, f"\n{bullet}")

    # 4) 最后，把 A． B． C． D． 放在新行
    pattern = r'(A．.*B．.*C．.*D．.*)'
    match = re.search(pattern, choice_text)
    if match:
        choices_str = match.group(1).strip()
        start_idx = match.start()
        end_idx = match.end()
        text_before = choice_text[:start_idx].strip()
        text_after = choice_text[end_idx:].strip()
        choice_text = f"{text_before}\n{choices_str}"
        if text_after:
            choice_text += f" {text_after}"

    # 5) 去除多余空白行
    lines = [x.rstrip() for x in choice_text.split('\n')]
    return lines

# 清理和拆分问题文本，分离题干、序号部分和选项部分
def clean_and_split_question(question_text):
    if not question_text:
        return "", [], ""
    
    question_text = str(question_text).strip()
    
    # 清理HTML实体
    replacements = {
        "&ldquo;": """,
        "&rdquo;": """,
        "&hellip;": "...",
        "&nbsp;": " ",
        "&mdash;": "—",
        "&rarr;": "",
        "&lsquo;": "'",
        "&rsquo;": "'",
        "&middot;": "",
        "&bull;": "",
        "&amp;": "&"
    }
    for old, new in replacements.items():
        question_text = question_text.replace(old, new)
    
    # 首先移除题目前面的原始题号（如 "3." 等）
    question_text = re.sub(r'^\s*\d+[．.、]\s*', '', question_text)
    
    # 使用正则表达式寻找序号部分（如①②③）
    bullet_pattern = r'(?m)^[①-⑳].*?(?=^[①-⑳]|$)'
    bullets = re.findall(bullet_pattern, question_text, re.MULTILINE)
    
    # 寻找选项部分（如A.xxx B.xxx）
    choice_pattern = r'(?m)(^[A-D][.．].*?(?=^[A-D][.．]|$))'
    choices = re.findall(choice_pattern, question_text, re.MULTILINE)
    
    # 提取选项部分的内容
    choice_part = '\n'.join(choices) if choices else ""
    
    # 如果找到了序号或选项，则移除题干中的这些部分
    if bullets:
        for bullet in bullets:
            question_text = question_text.replace(bullet, "")
    
    if choice_part:
        question_text = question_text.replace(choice_part, "")
    
    # 清理剩余部分作为题干
    questionPart = question_text.strip()
    
    return questionPart, bullets, choice_part

app = Flask(__name__, static_folder='static', template_folder='templates')
app.config.update(
    SECRET_KEY='生成一个随机的密钥',  # 设置一个固定的密钥
    WTF_CSRF_CHECK_DEFAULT=False,       # 关闭默认的CSRF检查
    WTF_CSRF_ENABLED=False             # 完全禁用CSRF保护
)
csrf = CSRFProtect(app)

# 确保instance目录存在
instance_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'instance')
if not os.path.exists(instance_path):
    os.makedirs(instance_path)

# 使用本地数据库文件
db_path = os.path.join(instance_path, 'xlct12.db')
app.config['SQLALCHEMY_DATABASE_URI'] = f'sqlite:///{os.path.abspath(db_path)}'
app.config['SQLALCHEMY_ENGINE_OPTIONS'] = {
    'connect_args': {
        'timeout': 30,
        'uri': True  # 启用更严格的URI解析
    },
    'pool_pre_ping': True  # 增加连接健康检查
}
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
app.config['SQLALCHEMY_COMMIT_ON_TEARDOWN'] = False  # 禁用自动提交
app.config['UPLOAD_FOLDER'] = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'uploads', 'papers')
if not os.path.exists(app.config['UPLOAD_FOLDER']):
    os.makedirs(app.config['UPLOAD_FOLDER'])

# 允许的文件扩展名
ALLOWED_EXTENSIONS = {'pdf', 'doc', 'docx', 'zip', 'rar'}

def allowed_file(filename):
    try:
        if not filename:
            return False
            
        # 解码URL编码字符并处理特殊符号
        filename = unquote(filename).strip()
        filename = filename.split('\\')[-1]  # 处理Windows路径
        
        # 使用更健壮的正则表达式
        match = re.search(r'\.([a-zA-Z0-9]+)$', filename)
        if not match:
            return False
            
        extension = match.group(1).lower().strip()
        allowed_extensions = {'pdf', 'doc', 'docx', 'zip', 'rar'}
        
        print(f"[DEBUG] 原始文件名: {filename}")
        print(f"[DEBUG] 解析扩展名: {extension}")
        print(f"[DEBUG] 允许的扩展名: {allowed_extensions}")
        
        return extension in allowed_extensions
    except Exception as e:
        print(f"文件检查错误: {str(e)}")
        return False

db = SQLAlchemy(app)

# 在创建数据库路径后添加验证
print(f"数据库文件路径：{db_path}")  # 输出实际路径

# 在数据库配置后添加详细检查
print(f"当前工作目录：{os.getcwd()}")
db_dir = os.path.dirname(db_path)
print(f"数据库目录可写性：{os.access(db_dir, os.W_OK)}")
print(f"数据库文件是否存在：{os.path.exists(db_path)}")
print(f"数据库文件权限：{oct(os.stat(db_path).st_mode)[-3:] if os.path.exists(db_path) else '文件不存在'}")

# 在配置后添加SQLAlchemy日志
logging.basicConfig()
logging.getLogger('sqlalchemy.engine').setLevel(logging.INFO)

# 在app.py中添加路径验证
print(f"当前使用的数据库文件绝对路径：{os.path.abspath(db_path)}")

# 定义 Subject 表
class Subject(db.Model):
    __tablename__ = 'subjects'
    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(100), unique=True)

# 定义 Textbook 表
class Textbook(db.Model):
    __tablename__ = 'textbooks'
    id = db.Column(db.Integer, primary_key=True)
    subject_id = db.Column(db.Integer, db.ForeignKey('subjects.id'))
    name = db.Column(db.String(100))

# 定义 Unit 表
class Unit(db.Model):
    __tablename__ = 'units'
    id = db.Column(db.Integer, primary_key=True)
    textbook_id = db.Column(db.Integer, db.ForeignKey('textbooks.id'))
    name = db.Column(db.String(100))

# 定义 Lesson 表
class Lesson(db.Model):
    __tablename__ = 'lessons'
    id = db.Column(db.Integer, primary_key=True)
    unit_id = db.Column(db.Integer, db.ForeignKey('units.id'))
    name = db.Column(db.String(100))

# 定义 QuestionType 表
class QuestionType(db.Model):
    __tablename__ = 'question_types'
    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(50), unique=True)

# 定义 Question 表
class Question(db.Model):
    __tablename__ = 'questions'
    id = db.Column(db.Integer, primary_key=True)
    lesson_id = db.Column(db.Integer, db.ForeignKey('lessons.id'))
    type_id = db.Column(db.Integer, db.ForeignKey('question_types.id'))
    content = db.Column(db.Text)
    answer = db.Column(db.Text)
    explanation = db.Column(db.Text)
    options = db.relationship('Option', backref='question')

# 定义 Option 表
class Option(db.Model):
    __tablename__ = 'options'
    id = db.Column(db.Integer, primary_key=True)
    question_id = db.Column(db.Integer, db.ForeignKey('questions.id'))
    content = db.Column(db.Text)
    is_correct = db.Column(db.Boolean)

# 新增 SU 模型类（添加到其他模型定义之后）
class SU(db.Model):
    __tablename__ = 'su'
    id = db.Column(db.Integer, primary_key=True, autoincrement=True)
    subject = db.Column(db.String(100))
    textbook = db.Column(db.String(100))
    chapter = db.Column(db.String(100))
    unit = db.Column(db.String(100))
    lesson = db.Column(db.String(100))
    question = db.Column(db.Text)
    answer = db.Column(db.Text)
    question_image = db.Column(db.LargeBinary)
    answer_image = db.Column(db.LargeBinary)
    question_image_filename = db.Column(db.String(100))
    answer_image_filename = db.Column(db.String(100))
    question_type = db.Column(db.String(50))
    education_stage = db.Column(db.String(50), default='高中')  # 新增学段字段，默认为高中
    audio_file_path = db.Column(db.String(255))  # 新增音频文件路径字段
    audio_filename = db.Column(db.String(100))  # 新增音频文件名字段
    audio_content = db.Column(db.LargeBinary)  # 新增音频文件内容字段
    question_number = db.Column(db.Integer)  # 新增题号字段，用于关联MP3文件

# 在其他模型定义后添加
class Paper(db.Model):
    __tablename__ = 'papers'
    
    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(200), nullable=False)  # 试卷名称
    region = db.Column(db.String(50), nullable=False)  # 地区
    subject = db.Column(db.String(50), nullable=False)  # 科目
    stage = db.Column(db.String(20), nullable=False)   # 阶段
    source = db.Column(db.String(100), nullable=False) # 来源
    source_type = db.Column(db.String(20), default='地区联考')  # 来源类型：地区联考、名校调考
    year = db.Column(db.Integer, nullable=False)       # 年份
    file_path = db.Column(db.String(500), nullable=False) # 文件路径
    upload_time = db.Column(db.DateTime, default=datetime.now) # 上传时间

    def to_dict(self):
        return {
            'id': self.id,
            'name': self.name,
            'region': self.region,
            'subject': self.subject,
            'stage': self.stage,
            'source': self.source,
            'source_type': self.source_type,
            'year': self.year,
            'file_path': self.file_path,
            'upload_time': self.upload_time.strftime('%Y-%m-%d %H:%M:%S')
        }

# 定义下载日志表
class DownloadLog(db.Model):
    __tablename__ = 'download_logs'
    
    id = db.Column(db.Integer, primary_key=True)
    paper_id = db.Column(db.Integer, nullable=False)  # 试卷ID
    paper_name = db.Column(db.String(200), nullable=False)  # 试卷名称
    file_path = db.Column(db.String(500), nullable=False)  # 文件路径
    user_ip = db.Column(db.String(50))  # 用户IP
    download_time = db.Column(db.DateTime, default=datetime.now)  # 下载时间

# 修改数据库初始化部分
with app.app_context():
    try:
        # 检查 papers 表是否存在
        result = db.engine.execute(text("SELECT name FROM sqlite_master WHERE type='table' AND name='papers'"))
        if not result.fetchone():
            print("创建 papers 表...")
            db.create_all()
        else:
            print("papers 表已存在")
        
        # 检查su表中是否有audio_file_path列
        try:
            conn = db.engine.connect()
            # 尝试查询表结构
            result = conn.execute(text("PRAGMA table_info(su)"))
            columns = [row[1] for row in result.fetchall()]
            
            # 检查是否需要添加音频相关字段
            if 'audio_file_path' not in columns:
                print("为su表添加audio_file_path字段...")
                conn.execute(text("ALTER TABLE su ADD COLUMN audio_file_path VARCHAR(255)"))
            
            if 'audio_filename' not in columns:
                print("为su表添加audio_filename字段...")
                conn.execute(text("ALTER TABLE su ADD COLUMN audio_filename VARCHAR(100)"))
            
            if 'audio_content' not in columns:
                print("为su表添加audio_content字段...")
                conn.execute(text("ALTER TABLE su ADD COLUMN audio_content BLOB"))
            
            if 'question_number' not in columns:
                print("为su表添加question_number字段...")
                conn.execute(text("ALTER TABLE su ADD COLUMN question_number INTEGER"))
                
            print("数据库结构更新完成")
        except Exception as e:
            print(f"更新数据库结构时出错: {str(e)}")
            traceback.print_exc()
        
        # 验证表结构
        papers_count = Paper.query.count()
        print(f"当前试卷数量: {papers_count}")
        
    except Exception as e:
        print(f"数据库初始化错误: {str(e)}")
        db.create_all()  # 尝试重新创建所有表
    
    print("数据库表初始化完成")

# 在应用启动时创建测试文件到instance目录
with open(os.path.join(instance_path, 'test.txt'), 'w') as f:
    f.write('test')
print(f"测试文件写入{'成功' if os.path.exists(os.path.join(instance_path, 'test.txt')) else '失败'}")

@app.before_request
def check_database_initialization():
    if not hasattr(app, 'db_initialized'):
        try:
            with db.engine.connect() as conn:
                # 检查 su 表
                result = conn.execute(text("SELECT name FROM sqlite_master WHERE type='table' AND name='su'"))
                if not result.fetchone():
                    raise RuntimeError("su表缺失")
                    
            app.db_initialized = True
            print("数据库健康检查通过")
            
        except Exception as e:
            print(f"数据库初始化验证失败：{str(e)}")
            raise

@app.route('/')
def index():
    return render_template('index.html', active_page='index')  # 渲染前端页面

def parse_questions(text):
    """解析题目文本，返回题目列表"""
    questions = []
    # 清理HTML标签
    clean_text = re.sub(r'<[^>]+>', '', text)
    
    # 匹配主题目，忽略子问题编号
    pattern = r'(?:^|\n)\s*(\d+)[．.、]\s*(.*?)(?=(?:\n\s*\d+[．.、])|$)'
    matches = re.finditer(pattern, clean_text, re.DOTALL)
    
    for match in matches:
        question_num = match.group(1)
        question_content = match.group(2).strip()
        # 保留子问题编号在内容中
        questions.append(f"{question_num}．{question_content}")
    
    app.logger.info(f"解析题目结果：\n题号: {[q.split('．')[0] for q in questions]}")
    app.logger.info(f"题目数量: {len(questions)}")
    return questions

def parse_answers(text):
    """解析答案文本，返回答案列表"""
    answers = []
    # 清理HTML标签
    clean_text = re.sub(r'<[^>]+>', '', text)
    
    # 匹配答案，忽略子问题编号
    pattern = r'(?:^|\n)\s*(\d+)[．.、]\s*(.*?)(?=(?:\n\s*\d+[．.、])|$)'
    matches = re.finditer(pattern, clean_text, re.DOTALL)
    
    for match in matches:
        answer_num = match.group(1)
        answer_content = match.group(2).strip()
        # 保留答案内容
        answers.append(answer_content)
    
    app.logger.info(f"解析答案数量: {len(answers)}")
    return answers

@app.route('/add_question', methods=['POST'])
def add_question():
    try:
        print("开始处理题目提交...")
        # 获取表单字段
        subject = request.form.get('subject')
        textbook = request.form.get('textbook')
        chapter = request.form.get('chapter')
        unit = request.form.get('unit')
        lesson = request.form.get('lesson')
        question_text = request.form.get('question_text')
        answer_text = request.form.get('answer_and_explanation')
        question_type = request.form.get('question_type')
        education_stage = request.form.get('education_stage', '高中')  # 获取学段，默认为高中
        
        # 解析题目时处理图片
        def process_images(content, prefix):
            images = {}
            pattern = r'\[img:(\d+)\]'
            matches = re.findall(pattern, content)
            for img_id in matches:
                file = request.files.get(f'{prefix}_image_{img_id}')
                if file and allowed_file(file.filename):
                    image_data = file.read()
                    filename = secure_filename(file.filename)
                    images[img_id] = (image_data, filename)
            return re.sub(pattern, '[图片]', content), images

        # 处理MP3文件上传
        def process_mp3_files():
            mp3_files = {}
            
            # 处理多文件上传字段
            if 'mp3_files' in request.files:
                files = request.files.getlist('mp3_files')
                for file in files:
                    if file and file.filename.endswith('.mp3'):
                        # 从文件名中提取题号
                        filename = secure_filename(file.filename)
                        # 尝试从文件名中提取数字作为题号
                        question_number = None
                        matches = re.findall(r'(\d+)', filename)
                        if matches:
                            question_number = int(matches[0])
                        
                        if question_number:
                            # 生成唯一文件名并保存文件
                            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
                            mp3_filename = f"{timestamp}_{question_number}_{filename}"
                            
                            # 确保音频目录存在
                            mp3_dir = os.path.join(app.config['UPLOAD_FOLDER'], 'audio')
                            os.makedirs(mp3_dir, exist_ok=True)
                            
                            mp3_path = os.path.join(mp3_dir, mp3_filename)
                            
                            # 保存文件
                            file.save(mp3_path)
                            
                            # 存储相对路径
                            project_root = os.path.dirname(os.path.abspath(__file__))
                            relative_path = os.path.relpath(mp3_path, project_root)
                            
                            mp3_files[question_number] = (relative_path, mp3_filename)
                            print(f"上传MP3文件 {filename} 关联到题号 {question_number}")
            
            # 处理单个文件上传字段（保持兼容性）
            mp3_count = 0
            while True:
                mp3_key = f'mp3_file_{mp3_count}'
                if mp3_key not in request.files:
                    break
                    
                file = request.files.get(mp3_key)
                if file and file.filename.endswith('.mp3'):
                    # 从文件名中提取题号
                    filename = secure_filename(file.filename)
                    # 尝试从文件名中提取数字作为题号
                    question_number = None
                    matches = re.findall(r'(\d+)', filename)
                    if matches:
                        question_number = int(matches[0])
                    
                    if question_number:
                        # 生成唯一文件名并保存文件
                        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
                        mp3_filename = f"{timestamp}_{question_number}_{filename}"
                        mp3_path = os.path.join(app.config['UPLOAD_FOLDER'], 'audio', mp3_filename)
                        
                        # 确保音频目录存在
                        os.makedirs(os.path.dirname(mp3_path), exist_ok=True)
                        
                        # 保存文件
                        file.save(mp3_path)
                        
                        # 存储相对路径
                        project_root = os.path.dirname(os.path.abspath(__file__))
                        relative_path = os.path.relpath(mp3_path, project_root)
                        
                        mp3_files[question_number] = (relative_path, mp3_filename)
                        print(f"上传MP3文件 {filename} 关联到题号 {question_number}")
                
                mp3_count += 1
                
            return mp3_files

        # 处理题目中的图片
        processed_questions, question_images = process_images(question_text, 'question')
        # 处理答案中的图片
        processed_answers, answer_images = process_images(answer_text, 'answer')
        # 处理MP3文件
        mp3_files = process_mp3_files()

        print(f"解析题目文本，长度：{len(processed_questions)}")
        questions = parse_questions(processed_questions)
        print(f"成功解析题目数量：{len(questions)}")
        
        print(f"解析答案文本，长度：{len(processed_answers)}")
        answers = parse_answers(processed_answers)
        app.logger.info(f"解析到的答案列表: {answers[:3]}")  # 打印前三个答案示例
        app.logger.info(f"成功解析答案数量：{len(answers)}")
        
        if len(questions) != len(answers):
            app.logger.error(f"题目数量({len(questions)})和答案数量({len(answers)})不匹配")
            return "题目和答案数量不匹配", 400
        
        # 批量创建题目记录
        for idx, (q, a) in enumerate(zip(questions, answers), start=1):
            # 获取当前题目的图片
            q_img_data, q_img_name = question_images.get(str(idx), (None, None))
            a_img_data, a_img_name = answer_images.get(str(idx), (None, None))
            
            # 获取当前题目的MP3文件（如果有）
            mp3_path, mp3_filename = mp3_files.get(idx, (None, None))
            
            print(f"正在创建题目记录：{q[:50]}...")
            new_su = SU(
                subject=subject,
                textbook=textbook,
                chapter=chapter,
                unit=unit,
                lesson=lesson,
                question_type=question_type,
                question=q,
                answer=a,
                education_stage=education_stage,
                question_image=q_img_data,
                answer_image=a_img_data,
                question_image_filename=q_img_name,
                answer_image_filename=a_img_name,
                audio_file_path=mp3_path,
                audio_filename=mp3_filename,
                question_number=idx  # 设置题号
            )
            db.session.add(new_su)
        
        print(f"共处理题目图片：{len(question_images)}张")
        print(f"共处理答案图片：{len(answer_images)}张")
        print(f"共处理MP3文件：{len(mp3_files)}个")
        
        print("提交数据库事务...")
        db.session.commit()
        print("数据库写入成功")
        app.logger.info(f"题目-答案对应关系：\n题目: {questions[:3]}\n答案: {answers[:3]}")
        return redirect('/show_su')
    except Exception as e:
        db.session.rollback()
        print(f"添加题目时出错: {str(e)}")
        print(f"错误详情: {traceback.format_exc()}")
        return f"添加题目失败: {str(e)}", 500

# 添加获取单个题目的路由（在/add_question路由之后添加）
@app.route('/questions/<int:id>')
def get_question(id):
    # 尝试从SU表中获取题目
    question = SU.query.get_or_404(id)
    
    # 创建一个包含题目详细信息的HTML页面
    return render_template('question_detail.html', 
                          question=question,
                          active_page='search')

@app.route('/test_db')
def test_db():
    try:
        with db.engine.connect() as conn:
            conn.execute(text("CREATE TABLE IF NOT EXISTS test (id INTEGER)"))
            conn.execute(text("INSERT INTO test VALUES (1)"))
            return '写入成功'
    except Exception as e:
        return f'写入失败: {str(e)}'

@app.route('/test_add')
def test_add():
    """测试直接添加数据到数据库"""
    try:
        # 创建测试题目
        q = Question(content="测试题目", type_id="Single Choice")
        db.session.add(q)
        db.session.flush()
        
        # 添加选项
        o1 = Option(question_id=q.id, content="正确选项", is_correct=True)
        o2 = Option(question_id=q.id, content="错误选项", is_correct=False)
        db.session.add_all([o1, o2])
        
        db.session.commit()
        return jsonify({
            'question_id': q.id,
            'options': [o1.id, o2.id]
        })
    except Exception as e:
        db.session.rollback()
        return f"测试失败：{str(e)}", 500

@app.route('/questions')
def get_all_questions():
    questions = Question.query.all()
    return jsonify([{
        'id': q.id,
        'content': q.content,
        'type_id': q.type_id
    } for q in questions])

# 修改连接测试部分
with app.app_context():
    try:
        with db.engine.connect() as conn:
            conn.execute(text("SELECT 1"))
        print("数据库连接成功")
    except Exception as e:
        print(f"数据库连接失败：{str(e)}")

@app.route('/upload')
def upload_page():
    # 获取所有学科
    subjects = Subject.query.all()
    # 获取所有题型
    question_types = QuestionType.query.all()
    
    # 添加 CKEditor 配置
    ckeditor_config = {
        'filebrowserUploadUrl': url_for('upload_image'),
        'filebrowserBrowseUrl': url_for('browse_images'),
        'imageUploadUrl': url_for('upload_image')
    }
    
    return render_template('upload.html', subjects=subjects, question_types=question_types, 
                          active_page='upload', ckeditor_config=ckeditor_config)

@app.route('/questions')
def show_questions():
    questions = Question.query.all()
    return render_template('questions.html', questions=questions)

@app.route('/check_data')
def check_data():
    questions = Question.query.options(db.joinedload(Question.options)).all()
    result = []
    for q in questions:
        result.append({
            '题目ID': q.id,
            '题干': q.content[:50],
            '正确答案': q.answer,
            '选项': [{'内容': o.content, '是否正确': o.is_correct} for o in q.options]
        })
    return jsonify(result)

@app.route('/show_su')
def show_su():
    # 查询所有 su 表中的数据
    records = SU.query.all()
    return render_template('su_questions.html', records=records)

@app.route('/get_image/<int:id>/<type>')
def get_image(id, type):
    try:
        record = SU.query.get_or_404(id)
        def get_mime_type(filename):
            if filename:
                ext = filename.rsplit('.', 1)[1].lower()
                return f'image/{ext}'
            return 'image/jpeg'  # 默认类型
        
        if type == 'question' and record.question_image:
            mime_type = get_mime_type(record.question_image_filename)
            return send_file(
                io.BytesIO(record.question_image),
                mimetype=mime_type,
                as_attachment=False,
                download_name=record.question_image_filename or 'question.jpg'
            )
        elif type == 'answer' and record.answer_image:
            mime_type = get_mime_type(record.answer_image_filename)
            return send_file(
                io.BytesIO(record.answer_image),
                mimetype=mime_type,
                as_attachment=False,
                download_name=record.answer_image_filename or 'answer.jpg'
            )
        return '图片不存在', 404
    except Exception as e:
        print(f"获取图片失败：{str(e)}")
        return '获取图片失败', 500

# 修改上传目录配置
UPLOAD_FOLDER = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'static', 'uploads')
if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)
    
ALLOWED_EXTENSIONS = {'png', 'jpg', 'jpeg', 'gif'}

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

@app.route('/upload_image', methods=['POST'])
def upload_image():
    try:
        if 'upload' not in request.files:
            print("未检测到上传文件")
            return """
            <script>
                window.parent.CKEDITOR.tools.callFunction({}, '', '请选择要上传的图片');
            </script>
            """.format(request.args.get('CKEditorFuncNum'))

        file = request.files['upload']
        if file.filename == '':
            print("文件名为空")
            return """
            <script>
                window.parent.CKEDITOR.tools.callFunction({}, '', '未选择文件');
            </script>
            """.format(request.args.get('CKEditorFuncNum'))

        if not allowed_file(file.filename):
            print(f"不支持的文件类型：{file.filename}")
            return """
            <script>
                window.parent.CKEDITOR.tools.callFunction({}, '', '仅支持 {} 格式');
            </script>
            """.format(request.args.get('CKEditorFuncNum'), ', '.join(ALLOWED_EXTENSIONS))

        # 生成安全的文件名
        filename = secure_filename(str(uuid.uuid4()) + os.path.splitext(file.filename)[1])
        file_path = os.path.abspath(os.path.join(app.config['UPLOAD_FOLDER'], filename))
        os.makedirs(os.path.dirname(file_path), exist_ok=True)  # 确保目录存在
        
        # 保存文件
        file.save(file_path)
        print(f"文件保存路径: {file_path}")
        print(f"路径是否存在: {os.path.exists(file_path)}")
        
        # 获取文件URL
        file_url = url_for('static', filename=f'uploads/{filename}', _external=True)
        print(f"文件访问URL：{file_url}")
        
        # 返回CKEditor期望的响应格式
        return """
        <script>
            window.parent.CKEDITOR.tools.callFunction({}, '{}');
        </script>
        """.format(request.args.get('CKEditorFuncNum'), file_url)
        
    except Exception as e:
        print(f"图片上传失败：{str(e)}")
        print(f"错误详情：{traceback.format_exc()}")
        return """
        <script>
            window.parent.CKEDITOR.tools.callFunction({}, '', '上传失败：{}');
        </script>
        """.format(request.args.get('CKEditorFuncNum'), str(e))

# 在 app.py 中添加一个新的路由来处理图片浏览
@app.route('/browse_images')
def browse_images():
    # 获取 uploads 目录下的所有图片
    files = []
    for filename in os.listdir(UPLOAD_FOLDER):
        if allowed_file(filename):
            files.append({
                'image': url_for('static', filename=f'uploads/{filename}'),
                'thumb': url_for('static', filename=f'uploads/{filename}'),
                'title': filename
            })
    return jsonify(files)

@app.route('/client')
def client():
    return render_template('client.html', active_page='client')

@app.route('/api/questions')
def api_questions():
    # 从SU表中获取所有题目
    questions = SU.query.all()
    result = []
    
    for q in questions:
        # 添加ID到每个问题
        question_dict = {
            "id": q.id,
            "subject": q.subject,
            "textbook": q.textbook,
            "chapter": q.chapter,
            "unit": q.unit,
            "lesson": q.lesson,
            "question": q.question,
            "answer": q.answer,
            "questionType": q.question_type,
            "educationStage": getattr(q, 'education_stage', '高中')  # 默认为高中
        }
        
        # 检查是否有问题图片和答案图片
        if q.question_image is not None and len(q.question_image) > 0:
            question_dict["has_question_image"] = True
            question_dict["question_image_url"] = f"/get_image/{q.id}/question"
        else:
            question_dict["has_question_image"] = False
            
        if q.answer_image is not None and len(q.answer_image) > 0:
            question_dict["has_answer_image"] = True
            question_dict["answer_image_url"] = f"/get_image/{q.id}/answer"
        else:
            question_dict["has_answer_image"] = False
        
        # 检查是否有音频文件
        if q.audio_file_path and q.audio_filename:
            question_dict["has_audio"] = True
            question_dict["audio_url"] = f"/get_audio/{q.id}"
            question_dict["question_number"] = q.question_number
        else:
            question_dict["has_audio"] = False
        
        result.append(question_dict)
        
    return jsonify(result)

@app.route('/clear_database', methods=['POST'])
@csrf.exempt
def clear_database():
    return "此功能已禁用", 403

@app.route('/generate_paper', methods=['POST'])
def generate_paper():
    try:
        print("Starting paper generation process...")
        # 1) 获取选中的题目 ID 和试卷标题
        data = request.get_json()
        if not data:
            print("Error: No JSON data received")
            return jsonify({'error': 'No data received'}), 400
            
        question_ids = data.get('question_ids', [])
        paper_title = data.get('paper_title', '试卷')
        
        if not question_ids:
            print("Error: No question IDs provided")
            return jsonify({'error': 'No question IDs provided'}), 400
        
        print(f"Received request to generate paper with title '{paper_title}' and {len(question_ids)} questions")
        
        # 查询所有题目
        try:
            all_questions = SU.query.filter(SU.id.in_(question_ids)).all()
            print(f"Retrieved {len(all_questions)} questions from database")
            
            if len(all_questions) == 0:
                print("Error: No questions found with the provided IDs")
                return jsonify({'error': 'No questions found with the provided IDs'}), 404
                
            # 创建ID到题目的映射
            question_map = {q.id: q for q in all_questions}
            
            # 按照原始顺序排列题目
            questions = [question_map[qid] for qid in question_ids if qid in question_map]
            print(f"Final question count for paper: {len(questions)}")
        except Exception as db_error:
            print(f"Database error: {str(db_error)}")
            traceback.print_exc()
            return jsonify({'error': f'Database error: {str(db_error)}'}), 500

        # 2) 创建 Word 文档
        try:
            doc = Document()
            
            # 设置A4纸张和标准边距
            sections = doc.sections
            for section in sections:
                # A4纸张尺寸 (8.27 x 11.69 英寸)
                section.page_width = Inches(8.27)
                section.page_height = Inches(11.69)
                # 设置标准边距
                section.left_margin = Inches(1.0)
                section.right_margin = Inches(1.0)
                section.top_margin = Inches(1.0)
                section.bottom_margin = Inches(1.0)
            
            # 设置基本样式
            style = doc.styles['Normal']
            style.font.name = '宋体'  # 使用宋体作为默认字体
            style.font.size = Pt(10.5)  # 10.5磅字号
            style.font.color.rgb = RGBColor(0, 0, 0)  # 黑色文字
            style.paragraph_format.line_spacing = 1.5  # 1.5倍行距
            style.paragraph_format.space_after = Pt(0)  # 默认段落间距
            
            # 创建标题样式（宋体，15磅，加粗，居中）
            title_style = doc.styles.add_style('Title Bold', WD_STYLE_TYPE.PARAGRAPH)
            title_style.base_style = doc.styles['Normal']
            title_style.font.name = '宋体'
            title_style.font.bold = True
            title_style.font.size = Pt(15)
            title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_style.paragraph_format.space_before = Pt(0)
            title_style.paragraph_format.space_after = Pt(10)
            
            # 创建章节标题样式（宋体，10.5磅，加粗，左对齐）
            section_style = doc.styles.add_style('Section Title', WD_STYLE_TYPE.PARAGRAPH)
            section_style.base_style = doc.styles['Normal']
            section_style.font.name = '宋体'
            section_style.font.bold = True
            section_style.font.size = Pt(10.5)
            section_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            section_style.paragraph_format.space_before = Pt(12)
            section_style.paragraph_format.space_after = Pt(6)
            
            # 创建选项样式（与题目内容相同的字体和字号）
            option_style = doc.styles.add_style('Option Style', WD_STYLE_TYPE.PARAGRAPH)
            option_style.base_style = doc.styles['Normal']
            option_style.font.name = '宋体'
            option_style.font.size = Pt(10.5)
            option_style.paragraph_format.line_spacing = 1.5
            option_style.paragraph_format.space_after = Pt(0)
            
            # 选项符号样式（①②③④）
            option_marker_style = doc.styles.add_style('Option Marker', WD_STYLE_TYPE.CHARACTER)
            option_marker_style.font.name = '宋体'
            option_marker_style.font.bold = False
            option_marker_style.font.size = Pt(10.5)
            
            # ABCD选项样式（与题目内容相同的字体和字号）
            choice_style = doc.styles.add_style('Choice Style', WD_STYLE_TYPE.PARAGRAPH)
            choice_style.base_style = doc.styles['Normal']
            choice_style.font.name = '宋体'
            choice_style.font.size = Pt(10.5)
            choice_style.paragraph_format.line_spacing = 1.5
            
            # 表格样式
            table_style = doc.styles.add_style('Table Style', WD_STYLE_TYPE.PARAGRAPH)
            table_style.base_style = doc.styles['Normal']
            table_style.font.name = '宋体'
            table_style.font.size = Pt(10.5)
            table_style.font.bold = False
            
            # 添加标题（宋体，15磅，加粗，居中）
            title_paragraph = doc.add_paragraph(paper_title)
            title_paragraph.style = doc.styles['Title Bold']
            title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception as doc_setup_error:
            print(f"Error setting up document: {str(doc_setup_error)}")
            traceback.print_exc()
            return jsonify({'error': f'Error setting up document: {str(doc_setup_error)}'}), 500

        # 3) 对题目按类型分组并添加章节标题
        try:
            # 映射中文数字
            chinese_numbers = {
                1: '一',
                2: '二',
                3: '三',
                4: '四',
                5: '五',
                6: '六',
                7: '七',
                8: '八',
                9: '九',
                10: '十'
            }
            
            # 对题目按照题型进行分组
            grouped_questions = {}
            for q in questions:
                q_type = q.question_type or '未分类题目'
                if q_type not in grouped_questions:
                    grouped_questions[q_type] = []
                grouped_questions[q_type].append(q)
            
            # 按照特定顺序排列题型组
            question_type_order = [
                '单选题', '多选题', '判断题', '填空题', '解答题', 
                '主观题', '计算题', '论述题', '作文',
                '诗词鉴赏', '文言文阅读', '现代文阅读'
            ]
            
            # 根据优先级排序题型组
            ordered_types = []
            # 先添加已知顺序的题型
            for q_type in question_type_order:
                if q_type in grouped_questions:
                    ordered_types.append(q_type)
            # 再添加其他题型
            for q_type in grouped_questions:
                if q_type not in ordered_types:
                    ordered_types.append(q_type)
        except Exception as group_error:
            print(f"Error grouping questions: {str(group_error)}")
            traceback.print_exc()
            return jsonify({'error': f'Error grouping questions: {str(group_error)}'}), 500

        # 4) 添加题目到文档
        try:
            # 遍历每个题型分组添加题目
            overall_question_num = 1
            for section_num, q_type in enumerate(ordered_types, 1):
                # 添加章节标题（一、单选题）
                section_title = f'{chinese_numbers.get(section_num, str(section_num))}、{q_type}'
                section_para = doc.add_paragraph(section_title)
                section_para.style = doc.styles['Section Title']
                
                # 添加该题型的所有题目
                for q in grouped_questions[q_type]:
                    # 使用 clean_and_split_question() 来拆分题干
                    questionPart, bulletsList, choicePart = clean_and_split_question(q.question)
                    
                    # 获取选项部分（无论是否在答案中，都放到题目下面）
                    answer_text = q.answer
                    option_pattern = r'([A-D]\.[^\n]+)'  # 匹配选项
                    answer_options = re.findall(option_pattern, answer_text)
                
                    # 将题干和题号放在一行
                    p = doc.add_paragraph(style='Normal')
                    p.paragraph_format.line_spacing = 1.5
                    p.paragraph_format.space_after = Pt(6)
                    
                    # 题号
                    question_number = p.add_run(f"{overall_question_num}. ")
                    question_number.font.bold = False
                    question_number.font.size = Pt(10.5)
                    
                    # 增加整体题号计数
                    overall_question_num += 1
                
                    # 正文部分 - 宋体，10.5磅
                    question_text = p.add_run(questionPart)
                    question_text.font.size = Pt(10.5)
                    
                    # 写入序号部分（每个序号各自占一行）
                    for bullet in bulletsList:
                        p = doc.add_paragraph(style='Option Style')
                        
                        # 提取序号和文本
                        bullet_text = bullet.strip()
                        marker_match = re.match(r'([①-⑳])(.*)', bullet_text)
                        
                        if marker_match:
                            # 有序号（①②③④等）
                            marker, content = marker_match.groups()
                            
                            # 序号不加粗
                            marker_run = p.add_run(marker)
                            marker_run.font.bold = False
                            marker_run.font.size = Pt(10.5)
                            
                            # 添加间距（使用空格模拟）
                            p.add_run(" ")
                            
                            # 选项文本
                            content_run = p.add_run(content)
                            content_run.font.size = Pt(10.5)
                        else:
                            # 没有序号的情况，直接添加文本
                            text_run = p.add_run(bullet_text)
                            text_run.font.size = Pt(10.5)
                    
                    # 添加选项 - 优先使用问题中的选项
                    if choicePart:
                        # 提取每个选项并检查样式
                        options_text = optimize_choice_display(choicePart)
                        has_abcd_options = any(opt.strip().startswith(('A．', 'B．', 'C．', 'D．')) for opt in options_text)
                        
                        # 如果选项文本包含 A, B, C, D 开头的内容，则使用它
                        if has_abcd_options:
                            # 检查是否是引用序号类型（如 A．①③）
                            has_reference_options = any(re.search(r'[A-D]．[①②③④]', opt) for opt in options_text)
                            
                            if has_reference_options:
                                # 引用选项类型放在同一行
                                abcd_options = [opt.strip() for opt in options_text 
                                              if opt.strip().startswith(('A．', 'B．', 'C．', 'D．'))]
                                p = doc.add_paragraph()
                                p.paragraph_format.space_after = Pt(0)
                                
                                # 设置选项字体为宋体10.5磅
                                run = p.add_run(' '.join(abcd_options))
                                run.font.size = Pt(10.5)
                            else:
                                # 普通选项各自占一行
                                for option in options_text:
                                    if option.strip() and option.strip().startswith(('A．', 'B．', 'C．', 'D．')):
                                        p = doc.add_paragraph()
                                        p.paragraph_format.space_after = Pt(0)
                                        run = p.add_run(option.strip())
                                        run.font.size = Pt(10.5)
                        # 如果问题中没有ABCD选项，但答案中有，则使用答案中的
                        elif answer_options:
                            # 检查是否是引用序号类型
                            has_reference_pattern = any(re.search(r'[A-D]\.[①②③④]', opt) for opt in answer_options)
                            
                            if has_reference_pattern:
                                p = doc.add_paragraph()
                                p.paragraph_format.space_after = Pt(0)
                                run = p.add_run(' '.join(answer_options))
                                run.font.size = Pt(10.5)
                            else:
                                # 添加与选项区域的间距
                                doc.add_paragraph().paragraph_format.space_after = Pt(6)
                                
                                for option in answer_options:
                                    p = doc.add_paragraph(style='Choice Style')
                                    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                                
                                    # 提取选项字母和内容
                                    choice_match = re.match(r'([A-D]\.)(.*)//', option)
                                    if choice_match:
                                        letter, choice_content = choice_match.groups()
                                    
                                        # 选项字母
                                        letter_run = p.add_run(letter)
                                        letter_run.font.size = Pt(10.5)
                                        
                                        # 选项内容
                                        content_run = p.add_run(choice_content)
                                        content_run.font.size = Pt(10.5)
                                    else:
                                        # 没有匹配到模式，直接添加整行
                                        run = p.add_run(option)
                                        run.font.size = Pt(10.5)
                    
                    # 题目间空行
                    doc.add_paragraph()
        except Exception as add_questions_error:
            print(f"Error adding questions to document: {str(add_questions_error)}")
            traceback.print_exc()
            return jsonify({'error': f'Error adding questions to document: {str(add_questions_error)}'}), 500

        # 5) 添加答案部分
        try:
            # --- 答案部分 ---
            doc.add_page_break()
            
            # 添加参考答案标题（居中显示）
            answer_heading = doc.add_heading('参考答案', 0)
            answer_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 设置标题字体为宋体和大小
            for run in answer_heading.runs:
                run.font.name = '宋体'
                run.font.size = Pt(15)
                run.font.bold = True
            
            # 检查是否有选择题
            choice_questions = []
            question_indices = []
            choice_answers = []
            
            for i, q in enumerate(questions):
                answer_text = q.answer
                # 匹配字母答案（A、B、C、D）
                match = re.search(r'\b([A-D])\b', answer_text)
                if match:
                    choice_questions.append(q)
                    question_indices.append(i + 1)  # 保存原始题号
                    choice_answers.append(match.group(1))
            
            # 如果有选择题，我们才创建表格
            if choice_questions:
                # 计算需要多少行（每行10题）
                rows_needed = math.ceil(len(choice_questions) / 10)
                
                # 创建表格
                table = doc.add_table(rows=rows_needed * 2, cols=min(len(choice_questions), 10))
                table.style = 'Table Grid'
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                
                # 设置表格宽度
                table.autofit = False
                
                # 设置行高
                for row in table.rows:
                    row.height = Pt(28)  # 设置行高，确保足够的垂直空间
                    row.height_rule = 1  # 1表示固定高度
                
                # 填充表格内容
                current_cell_index = 0
                
                for q_idx in range(len(choice_questions)):
                    row_idx = (q_idx // 10) * 2
                    col_idx = q_idx % 10
                    
                    # 确保我们不会超出表格范围
                    if col_idx < len(table.columns):
                        # 题号单元格
                        num_cell = table.cell(row_idx, col_idx)
                        num_cell.text = str(question_indices[q_idx])
                        num_paragraph = num_cell.paragraphs[0]
                        num_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = num_paragraph.runs[0]
                        run.font.name = '宋体'
                        run.font.bold = True
                        run.font.size = Pt(10.5)
                        
                        # 设置单元格垂直居中
                        num_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                        
                        # 答案单元格
                        ans_cell = table.cell(row_idx + 1, col_idx)
                        ans_cell.text = choice_answers[q_idx]
                        ans_paragraph = ans_cell.paragraphs[0]
                        ans_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = ans_paragraph.runs[0]
                        run.font.name = '宋体'
                        run.font.size = Pt(10.5)
                        
                        # 设置单元格垂直居中
                        ans_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                
                # 设置表格边框
                set_cell_border(table)
            
            # 添加表格后的空白 - 减少空白量
            doc.add_paragraph().paragraph_format.space_after = Pt(16)
            
            # 添加答案与解析标题
            explanation_heading = doc.add_heading('答案与解析', 1)
            
            # 设置标题字体为宋体和大小
            for run in explanation_heading.runs:
                run.font.name = '宋体'
                run.font.size = Pt(15)
                run.font.bold = True
            
            # 添加每道题的解析，不留空白
            for i, q in enumerate(questions, 1):
                # 获取原始答案文本
                answer_text = q.answer.strip()
                
                # 清理HTML实体
                replacements = {
                    "&ldquo;": "\"",
                    "&rdquo;": "\"",
                    "&hellip;": "...",
                    "&nbsp;": " ",
                    "&mdash;": "—",
                    "&rarr;": "",
                    "&lsquo;": "'",
                    "&rsquo;": "'",
                    "&middot;": "",
                    "&bull;": "",
                    "&amp;": "&"
                }
                for old, new in replacements.items():
                    answer_text = answer_text.replace(old, new)
                
                # 使用正则表达式移除所有剩余的HTML实体
                answer_text = re.sub(r'&[a-zA-Z0-9#]+;', '', answer_text)
                
                # 移除答案开头的题号（如：9．）
                answer_text = re.sub(r'^\d+[.\uff0e]\s*', '', answer_text)
                
                # 提取答案选项（A-D）
                option_match = re.search(r'\b([A-D])\b', answer_text)
                letter_answer = ''
                detailed_explanation = answer_text
                
                # 确保主观题的答案文本也会被保留
                if option_match:
                    letter_answer = option_match.group(1)
                else:
                    # 主观题不需要字母答案，直接显示完整答案
                    pass  # 不做处理，让详解部分显示完整答案
                
                # 添加题号行
                p = doc.add_paragraph()
                p.style = 'Normal'
                p.paragraph_format.space_after = Pt(0)  # 移除段落后的空白
                run = p.add_run(f"第{i}题：")
                run.font.bold = True
                run.font.size = Pt(10.5)
                
                # 添加答案选项行（如果有）
                if letter_answer:
                    p = doc.add_paragraph()
                    p.style = 'Normal'
                    p.paragraph_format.space_after = Pt(0)  # 移除段落后的空白
                    p.add_run(letter_answer).font.size = Pt(10.5)
                
                # 添加详解部分
                if '【详解】' in detailed_explanation:
                    # 处理详解文本，移除任何空行
                    parts = detailed_explanation.split('【详解】', 1)
                    if len(parts) > 1:
                        explanation_text = parts[1]
                        # 移除所有多余空行
                        explanation_text = re.sub(r'\n\s*\n', '\n', explanation_text)
                        # 将所有换行符替换为空格
                        explanation_text = re.sub(r'\n+', ' ', explanation_text)
                        
                        # 创建新段落
                        p = doc.add_paragraph()
                        p.style = 'Normal'
                        p.paragraph_format.space_after = Pt(0)  # 完全移除段落后的空白
                        p.paragraph_format.line_spacing = 1.5  # 设置行距为1.5倍
                        
                        # 添加详解标签
                        tag_run = p.add_run('【详解】')
                        tag_run.font.name = '宋体'
                        tag_run.font.bold = True
                        tag_run.font.size = Pt(10.5)
                        
                        # 添加处理后的详解文本（无空行）
                        content_run = p.add_run(explanation_text)
                        content_run.font.name = '宋体'
                        content_run.font.size = Pt(10.5)
                    elif detailed_explanation.strip():
                        # 如果没有详解标签但有内容
                        # 移除所有多余空行
                        clean_text = re.sub(r'\n\s*\n', '\n', detailed_explanation)
                        # 将所有换行符替换为空格
                        clean_text = re.sub(r'\n+', ' ', clean_text)
                        
                        # 如果没有选项答案，这可能是主观题答案
                        # 直接显示所有文本作为答案
                        p = doc.add_paragraph()
                        p.style = 'Normal'
                        p.paragraph_format.space_after = Pt(0)  # 完全移除段落后的空白
                        
                        # 如果没有选项答案且没有详解标记，则不添加详解标记
                        if not letter_answer and '【详解】' not in clean_text:
                            # 直接显示答案文本作为主观题答案
                            content_run = p.add_run(clean_text)
                            content_run.font.name = '宋体'
                            content_run.font.size = Pt(10.5)
                        else:
                            # 如果有选项答案或者显示含详解标记，则添加标准详解标记
                            tag_run = p.add_run('【详解】')
                            tag_run.font.name = '宋体'
                            tag_run.font.bold = True
                            tag_run.font.size = Pt(10.5)
                            
                            content_run = p.add_run(clean_text)
                            content_run.font.name = '宋体'
                            content_run.font.size = Pt(10.5)
                else:
                    # 如果没有详解标签但有内容
                    # 移除所有多余空行
                    clean_text = re.sub(r'\n\s*\n', '\n', detailed_explanation)
                    # 将所有换行符替换为空格
                    clean_text = re.sub(r'\n+', ' ', clean_text)
                    
                    # 如果没有选项答案，这可能是主观题答案
                    # 直接显示所有文本作为答案
                    p = doc.add_paragraph()
                    p.style = 'Normal'
                    p.paragraph_format.space_after = Pt(0)  # 完全移除段落后的空白
                    
                    # 如果没有选项答案且没有详解标记，则不添加详解标记
                    if not letter_answer and '【详解】' not in clean_text:
                        # 直接显示答案文本作为主观题答案
                        content_run = p.add_run(clean_text)
                        content_run.font.name = '宋体'
                        content_run.font.size = Pt(10.5)
                    else:
                        # 如果有选项答案或者显示含详解标记，则添加标准详解标记
                        tag_run = p.add_run('【详解】')
                        tag_run.font.name = '宋体'
                        tag_run.font.bold = True
                        tag_run.font.size = Pt(10.5)
                        
                        content_run = p.add_run(clean_text)
                        content_run.font.name = '宋体'
                        content_run.font.size = Pt(10.5)
        except Exception as add_answers_error:
            print(f"Error adding answers to document: {str(add_answers_error)}")
            traceback.print_exc()
            return jsonify({'error': f'Error adding answers to document: {str(add_answers_error)}'}), 500

        # 6) 准备文档用于返回
        try:
            # 返回 Word
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            print("Successfully created Word document")
        except Exception as doc_save_error:
            print(f"Error saving document: {str(doc_save_error)}")
            traceback.print_exc()
            return jsonify({'error': f'Error saving document: {str(doc_save_error)}'}), 500

        # 7) 处理音频文件
        try:
            # 检查是否有包含音频文件的题目
            audio_files = []
            for q in questions:
                try:
                    # 检查音频文件路径是否存在
                    if hasattr(q, 'audio_file_path') and q.audio_file_path:
                        audio_path = q.audio_file_path
                        # 如果是相对路径，转换为绝对路径
                        if not os.path.isabs(audio_path):
                            audio_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), audio_path)
                        
                        print(f"Found audio file path: {audio_path}")
                        
                        if os.path.exists(audio_path):
                            with open(audio_path, 'rb') as audio_file:
                                audio_content = audio_file.read()
                                audio_files.append({
                                    'filename': q.audio_filename or f"audio_{q.id}.mp3",
                                    'content': audio_content
                                })
                                print(f"Successfully read audio file: {audio_path}, size: {len(audio_content)} bytes")
                        else:
                            print(f"Audio file does not exist: {audio_path}")
                except Exception as audio_err:
                    print(f"Error processing audio for question {q.id}: {str(audio_err)}")
                    traceback.print_exc()
                    # 继续处理下一个题目，忽略这个错误
                    continue

            print(f"Total audio files found: {len(audio_files)}")
        except Exception as audio_error:
            print(f"Error processing audio files: {str(audio_error)}")
            traceback.print_exc()
            # 即使音频处理失败，我们也继续并返回Word文档
            print("Continuing without audio files due to error")
            audio_files = []

        # 8) 创建ZIP文件（如果有音频）或返回Word文档
        try:
            # 如果有音频文件，创建ZIP压缩包
            if audio_files:
                try:
                    print("Creating ZIP archive with audio files...")
                    # 创建一个新的内存文件对象作为ZIP档案
                    zip_buffer = BytesIO()
                    
                    # 创建ZIP文件
                    with zipfile.ZipFile(zip_buffer, 'a', zipfile.ZIP_DEFLATED, False) as zip_file:
                        # 添加Word文档
                        print("Adding Word document to ZIP...")
                        zip_file.writestr(f'{paper_title}.docx', buffer.getvalue())
                        
                        # 添加所有音频文件
                        for audio in audio_files:
                            try:
                                filename = audio['filename']
                                content = audio['content']
                                if content:  # 确保内容存在
                                    print(f"Adding audio file to ZIP: {filename}, size: {len(content) if content else 0} bytes")
                                    zip_file.writestr(filename, content)
                            except Exception as file_err:
                                print(f"Error adding file {audio.get('filename', 'unknown')}: {str(file_err)}")
                                traceback.print_exc()
                                # 继续添加其他文件，忽略这个错误
                    
                    # 准备ZIP文件以供下载
                    zip_buffer.seek(0)
                    print("Sending ZIP file response...")
                    response = send_file(
                        zip_buffer,
                        as_attachment=True,
                        download_name=f'{paper_title}_with_audio.zip',
                        mimetype='application/zip'
                    )
                    
                    # 添加内容处置头以确保浏览器正确处理
                    response.headers['Content-Disposition'] = f'attachment; filename="{paper_title}_with_audio.zip"'
                    return response
                except Exception as zip_err:
                    print(f"Error creating ZIP file: {str(zip_err)}")
                    traceback.print_exc()
                    # 如果创建ZIP失败，退回到只返回Word文档
                    print("Falling back to Word document only...")
            
            # 如果没有音频文件，则只返回Word文档
            print("Sending Word document response...")
            response = send_file(
                buffer,
                as_attachment=True,
                download_name=f'{paper_title}.docx',
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
            response.headers['Content-Disposition'] = f'attachment; filename="{paper_title}.docx"'
            return response
        except Exception as final_error:
            print(f"Final error in paper generation: {str(final_error)}")
            traceback.print_exc()
            # 这是最后的回退，如果前面的所有尝试都失败了
            return jsonify({'error': f'Failed to generate paper: {str(final_error)}'}), 500
    except Exception as e:
        print(f"Error generating paper: {str(e)}")
        traceback.print_exc()  # 输出详细的错误堆栈信息
        return jsonify({'error': str(e), 'traceback': traceback.format_exc()}), 500

@app.route('/upload_paper', methods=['GET', 'POST'])
def upload_paper():
    if request.method == 'GET':
        # 获取所有地区、学科、学段等数据
        regions = db.session.query(Paper.region).distinct().all()
        subjects = db.session.query(Paper.subject).distinct().all()
        stages = db.session.query(Paper.stage).distinct().all()
        source_types = db.session.query(Paper.source_type).distinct().all()
        sources = db.session.query(Paper.source).distinct().all()
        
        return render_template('upload_paper.html', 
                              regions=[r[0] for r in regions if r[0]], 
                              subjects=[s[0] for s in subjects if s[0]], 
                              stages=[s[0] for s in stages if s[0]], 
                              source_types=[st[0] for st in source_types if st[0]], 
                              sources=[s[0] for s in sources if s[0]],
                              active_page='upload_paper')
    
    try:
        print(f"接收的表单数据: {request.form}")
        
        # 验证必填字段
        required_fields = ['region', 'stage', 'source', 'year']
        for field in required_fields:
            if not request.form.get(field):
                return jsonify({'error': f'缺少必填字段: {field}'})
        
        # 获取来源类型，默认为地区联考
        source_type = request.form.get('source_type', '地区联考')
        
        # 获取文件总数
        file_count = int(request.form.get('file_count', 0))
        if file_count == 0:
            return jsonify({'error': '没有上传文件'})
        
        # 允许的文件类型
        allowed_extensions = {'pdf', 'doc', 'docx', 'zip', 'rar'}
        
        # 保存所有上传的文件
        uploaded_papers = []
        
        for i in range(file_count):
            file_key = f'paper_file_{i}'
            name_key = f'paper_name_{i}'
            subject_key = f'paper_subject_{i}'
            
            if file_key not in request.files:
                continue
                
            file = request.files[file_key]
            if file.filename == '':
                continue
            
            # 验证文件类型
            file_ext = file.filename.rsplit('.', 1)[1].lower() if '.' in file.filename else ''
            if file_ext not in allowed_extensions:
                return jsonify({'error': f'不支持的文件类型: {file_ext}。允许的类型: {", ".join(allowed_extensions)}'})
            
            # 获取试卷名称和科目
            paper_name = request.form.get(name_key, '').strip()
            # 从试卷名称中移除"ks5u"字符串（不区分大小写）
            paper_name = re.sub(r'ks5u', '', paper_name, flags=re.IGNORECASE)
            
            subject = request.form.get(subject_key, '').strip()
            
            if not subject:
                return jsonify({'error': f'文件 {file.filename} 无法识别科目'})
            
            # 处理文件保存
            original_filename = file.filename.split('\\')[-1]  # 处理Windows路径
            
            # 清除特定模式，如【KS5U 高考】，并删除所有"ks5u"字符串（不区分大小写）
            clean_filename = re.sub(r'【KS5U\s+高考】|【KS5U】|【高考】|【真题】', '', original_filename)
            clean_filename = re.sub(r'ks5u', '', clean_filename, flags=re.IGNORECASE)
            
            # 如果清理后文件名为空，则使用原文件名
            if not clean_filename.strip():
                clean_filename = original_filename
                
            safe_filename = secure_filename(clean_filename)
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            # 为每个文件添加唯一时间戳，防止文件名冲突
            final_filename = f"{timestamp}_{i}_{safe_filename}"
            
            # 生成文件保存路径 - 直接保存在uploads/papers目录下，不再使用papers子目录
            file_path = os.path.abspath(os.path.join(
                app.config['UPLOAD_FOLDER'], 
                final_filename
            ))

            # 确保目录存在
            os.makedirs(os.path.dirname(file_path), exist_ok=True)
            
            # 保存文件
            file.save(file_path)
            
            # 创建数据库记录 - 存储相对于项目根目录的路径
            project_root = os.path.dirname(os.path.abspath(__file__))
            relative_path = os.path.relpath(file_path, project_root)
            
            new_paper = Paper(
                name=paper_name,
                region=request.form['region'],
                subject=subject,  # 使用每个文件独立的科目信息
                stage=request.form['stage'],
                source=request.form['source'],
                source_type=source_type,
                year=int(request.form['year']),
                file_path=relative_path  # 使用相对路径
            )
            
            db.session.add(new_paper)
            uploaded_papers.append(new_paper)
        
        # 提交所有新增的试卷
        db.session.commit()
        
        return jsonify({
            'success': True,
            'message': f'成功上传 {len(uploaded_papers)} 份试卷',
            'paper_ids': [paper.id for paper in uploaded_papers]
        })
        
    except Exception as e:
        print(f"上传试卷出错: {str(e)}")
        return jsonify({'error': f'上传失败: {str(e)}'})

@app.route('/download_paper/<int:paper_id>')
def download_paper(paper_id):
    try:
        paper = Paper.query.get_or_404(paper_id)
        file_path = paper.file_path
        
        # 提取文件名
        file_name = os.path.basename(file_path)
        
        # 获取项目根目录
        project_root = os.path.dirname(os.path.abspath(__file__))
        
        # 可能的文件路径列表
        possible_paths = [
            # 原始路径
            file_path,
            # 绝对路径
            os.path.join(project_root, file_path),
            # uploads/papers 目录
            os.path.join(project_root, 'uploads', 'papers', file_name),
            # 直接在 uploads 目录下
            os.path.join(project_root, 'uploads', file_name)
        ]
        
        # 移除None和重复项
        possible_paths = list(set([p for p in possible_paths if p]))
        
        # 记录尝试的路径
        print(f"尝试查找文件: {file_name}")
        for i, path in enumerate(possible_paths):
            print(f"路径{i+1}: {path}")
            if os.path.exists(path):
                print(f"文件找到于: {path}")
                return send_file(
                    path,
                    as_attachment=True,
                    download_name=paper.name + os.path.splitext(path)[1]
                )
        
        # 如果直接查找失败，尝试在uploads目录中递归搜索
        uploads_dir = os.path.join(project_root, 'uploads')
        print(f"在uploads目录中递归搜索: {uploads_dir}")
        
        if os.path.exists(uploads_dir):
            for root, dirs, files in os.walk(uploads_dir):
                for f in files:
                    if f == file_name:
                        found_path = os.path.join(root, f)
                        print(f"递归搜索找到文件: {found_path}")
                        return send_file(
                            found_path,
                            as_attachment=True,
                            download_name=paper.name + os.path.splitext(found_path)[1]
                        )
        
        # 如果所有尝试都失败，返回友好的错误页面
        print(f"错误：无法找到文件 {file_name} - 尝试路径: {', '.join(possible_paths)}")
        error_message = f"很抱歉，文件 '{paper.name}' 暂时无法下载。我们已记录此问题，请稍后再试或联系管理员。"
        return render_template('error.html', error=error_message, paper=paper), 404
        
    except Exception as e:
        print(f"下载文件时出错: {str(e)}")
        error_message = "下载过程中发生错误，请稍后再试或联系管理员。"
        return render_template('error.html', error=error_message), 500

@app.route('/admin/check_missing_files')
def check_missing_files():
    """检查缺失的文件"""
    try:
        # 获取所有试卷
        papers = Paper.query.all()
        missing_files = []
        
        for paper in papers:
            # 获取当前文件名
            file_name = os.path.basename(paper.file_path)
            
            # 查找文件的实际位置
            actual_path = find_actual_file_location(file_name)
            
            if not actual_path:
                # 文件缺失，添加到列表
                missing_files.append({
                    'id': paper.id,
                    'name': paper.name,
                    'subject': paper.subject,
                    'year': paper.year,
                    'file_path': paper.file_path,
                    'file_name': file_name
                })
        
        return jsonify({
            'success': True,
            'message': f'找到 {len(missing_files)} 个缺失文件',
            'missing_files': missing_files
        })
    except Exception as e:
        print(f"检查缺失文件时出错: {str(e)}")
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500

@app.route('/admin/replace_file', methods=['POST'])
def replace_file():
    """替换缺失的文件"""
    try:
        # 检查是否有文件上传
        if 'replacement_file' not in request.files:
            return jsonify({
                'success': False,
                'error': '没有上传文件'
            }), 400
        
        file = request.files['replacement_file']
        if file.filename == '':
            return jsonify({
                'success': False,
                'error': '未选择文件'
            }), 400
        
        # 检查文件类型
        allowed_extensions = {'pdf', 'doc', 'docx', 'zip', 'rar'}
        file_ext = file.filename.rsplit('.', 1)[1].lower() if '.' in file.filename else ''
        if file_ext not in allowed_extensions:
            return jsonify({
                'success': False,
                'error': f'不支持的文件类型，允许的类型: {", ".join(allowed_extensions)}'
            }), 400
        
        # 获取试卷ID
        paper_id = request.form.get('paper_id')
        if not paper_id:
            return jsonify({
                'success': False,
                'error': '缺少试卷ID'
            }), 400
        
        # 查找试卷
        paper = Paper.query.get(paper_id)
        if not paper:
            return jsonify({
                'success': False,
                'error': f'找不到ID为 {paper_id} 的试卷'
            }), 404
        
        # 生成新的文件名
        original_filename = os.path.basename(paper.file_path)
        file_ext = os.path.splitext(file.filename)[1]
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        new_filename = f"{timestamp}_{paper_id}{file_ext}"
        
        # 生成文件保存路径
        file_path = os.path.abspath(os.path.join(
            app.config['UPLOAD_FOLDER'], 
            new_filename
        ))

        # 确保目录存在
        os.makedirs(os.path.dirname(file_path), exist_ok=True)
        
        # 保存文件
        file.save(file_path)
        
        # 更新数据库中的文件路径
        project_root = os.path.dirname(os.path.abspath(__file__))
        relative_path = os.path.relpath(file_path, project_root)
        paper.file_path = relative_path
        
        # 提交更改
        db.session.commit()
        
        return jsonify({
            'success': True,
            'message': f'成功替换文件 "{paper.name}"',
            'new_path': relative_path
        })
    except Exception as e:
        print(f"替换文件时出错: {str(e)}")
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500

@app.route('/admin/fix_file_paths')
def fix_file_paths():
    """管理员路由：修复数据库中的文件路径"""
    try:
        # 获取所有试卷
        papers = Paper.query.all()
        fixed_count = 0
        not_found_count = 0
        
        for paper in papers:
            # 获取当前文件名
            file_name = os.path.basename(paper.file_path)
            
            # 查找文件的实际位置
            actual_path = find_actual_file_location(file_name)
            
            if actual_path:
                # 计算相对于项目根目录的路径
                project_root = os.path.dirname(os.path.abspath(__file__))
                if actual_path.startswith(project_root):
                    relative_path = os.path.relpath(actual_path, project_root)
                    
                    # 更新数据库中的路径
                    if paper.file_path != relative_path:
                        paper.file_path = relative_path
                        fixed_count += 1
                        print(f"已修复试卷 {paper.id}: {paper.name} 的路径为 {relative_path}")
            else:
                not_found_count += 1
                print(f"未找到试卷 {paper.id}: {paper.name} 的文件 {file_name}")
        
        # 提交更改
        if fixed_count > 0:
            db.session.commit()
            
        return jsonify({
            'success': True,
            'message': f'已修复 {fixed_count} 份试卷的路径，{not_found_count} 份试卷的文件未找到'
        })
    except Exception as e:
        print(f"修复文件路径时出错: {str(e)}")
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500

# 用于寻找文件实际位置的辅助函数 - 增强版
def find_actual_file_location(file_name, paper_info=None):
    """
    查找文件的实际位置
    
    Args:
        file_name: 文件名
        paper_info: 可选的试卷信息字典，包含id, name, year, subject, region等字段
    
    Returns:
        找到的文件路径，如果未找到则返回None
    """
    if not file_name:
        return None
    
    # 获取项目根目录
    project_root = os.path.dirname(os.path.abspath(__file__))
    
    # 可能的文件位置优先级列表
    search_dirs = [
        os.path.join(project_root, 'uploads', 'papers'),
        os.path.join(project_root, 'uploads', 'papers', 'papers'),
        os.path.join(project_root, 'uploads')
    ]
    
    # 1. 首先尝试直接匹配文件名
    for search_dir in search_dirs:
        if os.path.exists(search_dir):
            full_path = os.path.join(search_dir, file_name)
            if os.path.isfile(full_path):
                return full_path
    
    # 2. 如果直接匹配失败，尝试在所有目录中深度搜索
    all_files = []
    for search_dir in search_dirs:
        if os.path.exists(search_dir):
            for root, _, files in os.walk(search_dir):
                for f in files:
                    if f.endswith(('.pdf', '.doc', '.docx', '.zip', '.rar')):
                        all_files.append(os.path.join(root, f))
    
    # 如果提供了试卷信息，可以使用更高级的匹配算法
    if paper_info:
        # 优先使用ID匹配
        if 'id' in paper_info:
            id_pattern = f"_{paper_info['id']}_"
            id_matches = [f for f in all_files if id_pattern in os.path.basename(f)]
            if id_matches:
                return id_matches[0]
        
        # 使用关键词匹配
        if all(k in paper_info for k in ['year', 'subject', 'name']):
            year_str = str(paper_info['year'])
            subject = paper_info['subject'].lower()
            paper_name = paper_info['name'].lower()
            
            # 为"云学名校联盟"系列试卷设置特殊匹配逻辑
            if "云学名校联盟" in paper_name and "联考" in paper_name:
                for file_path in all_files:
                    file_basename = os.path.basename(file_path).lower()
                    
                    # 防止空元素错误
                    score = 0
                    
                    # 检查年份
                    if year_str in file_basename:
                        score += 10
                    
                    # 检查科目
                    if subject in file_basename:
                        score += 8
                    
                    # 检查地区
                    if 'region' in paper_info and paper_info['region'] and paper_info['region'].lower() in file_basename:
                        score += 7
                    
                    # 检查关键词
                    keywords = ["联考", "月", "上学期", "下学期"]
                    for keyword in keywords:
                        if keyword in paper_name and keyword in file_basename:
                            score += 3
                    
                    # 特定匹配项
                    if "云学" in file_basename or "名校联盟" in file_basename:
                        score += 5
                    
                    # 分数足够高则认为匹配成功
                    if score >= 15:
                        return file_path
            
            # 常规年份+科目匹配
            for file_path in all_files:
                file_basename = os.path.basename(file_path).lower()
                if year_str in file_basename and subject in file_basename:
                    # 如果还有地区信息，检查是否匹配
                    if 'region' in paper_info and paper_info['region']:
                        region = paper_info['region'].lower()
                        if region in file_basename:
                            return file_path
                    else:
                        return file_path
    
    # 3. 尝试基于部分文件名匹配
    file_base = os.path.splitext(file_name)[0]
    for file_path in all_files:
        if file_base in os.path.basename(file_path):
            return file_path
    
    # 4. 最后尝试模糊匹配
    # 移除数字前缀和特殊字符，提高匹配成功率
    clean_name = re.sub(r'^\d+_\d+_\d+_', '', file_base)
    clean_name = re.sub(r'[^\w\s\u4e00-\u9fff]+', '', clean_name).lower()
    
    if clean_name:
        for file_path in all_files:
            if clean_name in os.path.basename(file_path).lower():
                return file_path
    
    # 如果都失败了，返回None
    return None

@app.route('/admin')
def admin_page():
    """管理员页面"""
    return render_template('admin.html')

# 添加papers路由和函数
@app.route('/papers')
def papers_list():
    try:
        print("进入papers_list函数")
        
        # 获取分页和筛选参数
        page = request.args.get('page', 1, type=int)
        per_page = request.args.get('per_page', 20, type=int) # 默认每页20条
        region = request.args.get('region')
        subject = request.args.get('subject')
        stage = request.args.get('stage')
        source_type = request.args.get('source_type')
        source = request.args.get('source')
        year = request.args.get('year')
        keyword = request.args.get('keyword')
        
        # 获取所有试卷，使用复合排序
        papers_query = Paper.query
        
        # 应用筛选条件
        if region:
            papers_query = papers_query.filter(Paper.region == region)
        if subject:
            papers_query = papers_query.filter(Paper.subject == subject)
        if stage:
            papers_query = papers_query.filter(Paper.stage == stage)
        if source_type:
            papers_query = papers_query.filter(Paper.source_type == source_type)
        if source:
            papers_query = papers_query.filter(Paper.source == source)
        if year:
            try:
                if isinstance(year, str) and year.isdigit():
                    papers_query = papers_query.filter(Paper.year == int(year))
                elif isinstance(year, (int, float)):
                    papers_query = papers_query.filter(Paper.year == int(year))
            except (ValueError, TypeError) as e:
                print(f"年份转换错误: {e}")
        if keyword:
            papers_query = papers_query.filter(Paper.name.like(f'%{keyword}%'))
        
        print(f"应用筛选条件: 地区={region}, 科目={subject}, 学段={stage}, 类型={source_type}, 来源={source}, 年份={year}, 关键词={keyword}")
        
        try:
            # 使用复合排序条件：年份降序 > 湖北省优先 > 下学期优先 > 月份降序 > 上传时间降序
            sorted_query = papers_query.order_by(
                # 首先按照年份降序排序(较新的年份优先)
                Paper.year.desc(),
                # 其次湖北省的试卷优先显示
                db.case(
                    [(Paper.region == '湖北', 1)],
                    else_=0
                ).desc(),
                # 下学期优先于上学期
                db.case((Paper.name.like('%下学期%'), 2), (Paper.name.like('%上学期%'), 1), else_=0).desc(),
                # 3月优先于2月优先于1月
                db.case((Paper.name.like('%3月%'), 3), (Paper.name.like('%2月%'), 2), (Paper.name.like('%1月%'), 1), else_=0).desc(),
                # 最后按上传时间降序排序(最新上传的优先)
                Paper.upload_time.desc()
            )
            
            # 获取总试卷数
            total_papers = sorted_query.count()
            print(f"总共 {total_papers} 份试卷")
            
            # 实现分页
            paginated_papers = sorted_query.paginate(page=page, per_page=per_page, error_out=False)
            papers = paginated_papers.items
            
            print(f"成功查询到当前页 {len(papers)} 份试卷")
        except Exception as query_error:
            print(f"排序查询失败: {str(query_error)}")
            traceback.print_exc()
            # 使用简单查询作为备选
            sorted_query = papers_query.order_by(Paper.year.desc())
            total_papers = sorted_query.count()
            paginated_papers = sorted_query.paginate(page=page, per_page=per_page, error_out=False)
            papers = paginated_papers.items
            print(f"备选查询成功，获取到 {len(papers)} 份试卷")
        
        try:
            # 获取所有地区、学科、学段、来源类型和来源
            regions = db.session.query(Paper.region).distinct().all()
            subjects = db.session.query(Paper.subject).distinct().all()
            stages = db.session.query(Paper.stage).distinct().all()
            source_types = db.session.query(Paper.source_type).distinct().all()
            sources = db.session.query(Paper.source).distinct().all()
            print("成功获取筛选选项")
        except Exception as filter_error:
            print(f"获取筛选选项失败: {str(filter_error)}")
            traceback.print_exc()
            # 提供空列表作为备选
            regions = []
            subjects = []
            stages = []
            source_types = []
            sources = []
        
        # 准备渲染模板
        print("开始渲染模板")
        return render_template('papers.html', 
                             papers=papers, 
                             regions=regions, 
                             subjects=subjects, 
                             stages=stages, 
                             source_types=source_types, 
                             sources=sources, 
                             active_page='papers',
                             pagination=paginated_papers,
                             per_page=per_page,
                             page=page,
                             total=total_papers,
                             # 传递当前的筛选参数到模板
                             current_filters={
                                 'region': region,
                                 'subject': subject,
                                 'stage': stage,
                                 'source_type': source_type,
                                 'source': source,
                                 'year': year,
                                 'keyword': keyword
                             })
    except Exception as e:
        print(f"获取试卷列表失败: {str(e)}")
        traceback.print_exc()
        return render_template('papers.html', papers=[], error=str(e))

# 添加筛选功能路由
@app.route('/filter_papers', methods=['GET', 'POST'])
def filter_papers():
    try:
        if request.method == 'POST':
            # 从POST请求JSON中获取筛选参数
            filter_data = request.get_json()
            region = filter_data.get('region')
            subject = filter_data.get('subject')
            stage = filter_data.get('stage')
            source_type = filter_data.get('source_type')
            source = filter_data.get('source')
            year = filter_data.get('year')
            keyword = filter_data.get('keyword')
            page = filter_data.get('page', 1)
            per_page = filter_data.get('per_page', 20)
        else:  # GET请求
            # 从URL参数获取筛选条件
            region = request.args.get('region')
            subject = request.args.get('subject')
            stage = request.args.get('stage')
            source_type = request.args.get('source_type')
            source = request.args.get('source')
            year = request.args.get('year')
            keyword = request.args.get('keyword')
            page = request.args.get('page', 1, type=int)
            per_page = request.args.get('per_page', 20, type=int)
        
        print(f"筛选参数 [{request.method}]: 地区={region}, 科目={subject}, 学段={stage}, 类型={source_type}, 来源={source}, 年份={year}, 关键词={keyword}, 页码={page}")
        
        # 构建查询
        query = Paper.query
        
        # 添加筛选条件
        if region:
            query = query.filter(Paper.region == region)
        if subject:
            query = query.filter(Paper.subject == subject)
        if stage:
            query = query.filter(Paper.stage == stage)
        if source_type:
            query = query.filter(Paper.source_type == source_type)
        if source:
            query = query.filter(Paper.source == source)
        if year:
            try:
                if isinstance(year, str) and year.isdigit():
                    query = query.filter(Paper.year == int(year))
                elif isinstance(year, (int, float)):
                    query = query.filter(Paper.year == int(year))
            except (ValueError, TypeError) as e:
                print(f"年份转换错误: {e}")
        if keyword:
            query = query.filter(Paper.name.like(f'%{keyword}%'))
            
        # 排序逻辑
        sorted_query = query.order_by(
            Paper.year.desc(),
            db.case((Paper.region == '湖北', 1), else_=0).desc(),
            db.case((Paper.name.like('%下学期%'), 2), (Paper.name.like('%上学期%'), 1), else_=0).desc(),
            db.case((Paper.name.like('%3月%'), 3), (Paper.name.like('%2月%'), 2), (Paper.name.like('%1月%'), 1), else_=0).desc(),
            Paper.upload_time.desc()
        )
        
        # 计算总数
        total = sorted_query.count()
        print(f"筛选结果总数: {total}")
        
        # 分页
        paginated = sorted_query.paginate(page=page, per_page=per_page, error_out=False)
        papers = paginated.items
        
        # 构建试卷数据
        papers_data = []
        for paper in papers:
            try:
                if hasattr(paper, 'to_dict'):
                    papers_data.append(paper.to_dict())
                else:
                    # 手动构建字典
                    papers_data.append({
                        'id': paper.id,
                        'name': paper.name,
                        'region': paper.region,
                        'subject': paper.subject,
                        'stage': paper.stage,
                        'source_type': paper.source_type,
                        'source': paper.source,
                        'year': paper.year
                    })
            except Exception as paper_error:
                print(f"处理试卷数据时出错: {str(paper_error)}")
        
        return jsonify(papers_data)
    except Exception as e:
        print(f"筛选试卷出错: {str(e)}")
        traceback.print_exc()
        return jsonify({
            'error': str(e)
        }), 500

# 添加搜索功能路由
@app.route('/search')
def search():
    try:
        query = request.args.get('q', '')
        if not query or len(query) < 2:
            return render_template('search_results.html', 
                                  results=[], 
                                  query='', 
                                  count=0, 
                                  active_page='search')
        
        # 优化搜索性能：使用更精确的查询条件，避免过度使用通配符
        # 如果查询词很短，使用前缀匹配而不是全文匹配
        if len(query) <= 3:
            # 短查询使用前缀匹配
            questions = SU.query.filter(
                (SU.question.like(f'{query}%')) | 
                (SU.subject == query) |
                (SU.lesson == query)
            ).limit(50).all()
        else:
            # 较长查询使用包含匹配，但避免在所有字段上都使用
            questions = SU.query.filter(
                (SU.question.like(f'%{query}%')) | 
                (SU.answer.like(f'%{query}%')) |
                (SU.subject.like(f'%{query}%')) |
                (SU.lesson.like(f'%{query}%')) |
                (SU.unit.like(f'%{query}%'))
            ).limit(30).all()  # 减少结果数量以提高性能
        
        # 格式化结果
        results = []
        for q in questions:
            try:
                # 确保所有值都是基本类型且有效，避免序列化错误
                # 提取题目的前100个字符作为预览，确保无空值
                question_text = str(q.question or '')
                preview = question_text[:100] + '...' if len(question_text) > 100 else question_text
                
                # 将所有字段强制转换为基本类型，确保可序列化
                result_item = {
                    'id': int(q.id) if q.id is not None else 0,
                    'subject': str(q.subject or ''),
                    'unit': str(q.unit or ''),
                    'lesson': str(q.lesson or ''),
                    'preview': preview,  # 保留预览字段以兼容当前逻辑
                    'question': str(q.question or ''),  # 返回完整题目
                    'question_type': str(q.question_type or '')
                }
                
                # 确保没有undefined或特殊类型的字段
                for key, value in list(result_item.items()):
                    if value is None or type(value).__name__ not in ['str', 'int', 'float', 'bool', 'list', 'dict']:
                        result_item[key] = ''
                        
                results.append(result_item)
            except Exception as item_error:
                print(f"处理搜索结果项时出错: {str(item_error)}")
                # 跳过有问题的项，继续处理其他结果
        
        return render_template('search_results.html', 
                              results=results, 
                              query=query, 
                              count=len(results),
                              active_page='search')
                              
    except Exception as e:
        # 记录错误并返回友好的错误页面
        print(f"搜索时出错: {str(e)}")
        traceback.print_exc()
        return render_template('search_results.html',
                             results=[], 
                             query=query if 'query' in locals() else '', 
                             count=0,
                             error="搜索时出现错误，请稍后再试或修改搜索词。",
                             active_page='search')

@app.route('/get_audio/<int:id>')
def get_audio(id):
    try:
        record = SU.query.get_or_404(id)
        print(f"获取音频ID: {id}, 文件名: {record.audio_filename}, 路径: {record.audio_file_path}")
        
        # 检查和记录音频字段状态
        has_audio_content = hasattr(record, 'audio_content') and record.audio_content is not None and len(record.audio_content or b'') > 0
        has_audio_path = hasattr(record, 'audio_file_path') and record.audio_file_path is not None and len(record.audio_file_path or '') > 0
        
        print(f"音频状态: 内容存在={has_audio_content}, 路径存在={has_audio_path}")
        
        # 尝试从内容发送
        if has_audio_content:
            print(f"从二进制内容发送音频")
            return send_file(
                io.BytesIO(record.audio_content),
                mimetype='audio/mpeg',
                as_attachment=False,
                download_name=record.audio_filename or 'audio.mp3'
            )
        
        # 尝试从文件路径发送
        elif has_audio_path:
            audio_path = record.audio_file_path
            # 如果是相对路径，转换为绝对路径
            if not os.path.isabs(audio_path):
                audio_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), audio_path)
                
            print(f"音频绝对路径: {audio_path}")
            
            if os.path.exists(audio_path):
                print(f"从文件路径发送音频")
                return send_file(
                    audio_path,
                    mimetype='audio/mpeg',
                    as_attachment=False,
                    download_name=os.path.basename(audio_path)
                )
            else:
                print(f"音频文件路径存在但文件不存在: {audio_path}")
        
        # 无音频情况下返回一个静态MP3作为替代(防止404)
        default_audio_path = os.path.join(app.static_folder, 'audio', 'silence.mp3')
        
        # 如果默认音频存在就返回它
        if os.path.exists(default_audio_path):
            print(f"返回默认静音音频")
            return send_file(
                default_audio_path,
                mimetype='audio/mpeg',
                as_attachment=False,
                download_name='silence.mp3'
            )
        
        # 如果连默认音频都不存在，则生成临时的静音文件
        print(f"生成临时静音音频")
        silence = b'\xFF\xFB\x90\x04\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00'
        return send_file(
            io.BytesIO(silence),
            mimetype='audio/mpeg',
            as_attachment=False,
            download_name='silence.mp3'
        )
    except Exception as e:
        print(f"获取音频失败：{str(e)}")
        traceback.print_exc()
        
        # 即使出错也返回静音文件而不是错误，避免前端报错
        try:
            silence = b'\xFF\xFB\x90\x04\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00'
            return send_file(
                io.BytesIO(silence),
                mimetype='audio/mpeg',
                as_attachment=False,
                download_name='silence.mp3'
            )
        except:
            return '音频不可用', 200  # 返回200而不是错误状态码

@app.route('/logo.png')
def serve_logo():
    return send_file('logo.png')

@app.route('/fix_pagination.js')
def serve_fix_pagination_js():
    try:
        # 更改缓存控制标头，防止浏览器缓存旧版本
        response = send_from_directory(os.path.dirname(os.path.abspath(__file__)), 'fix_pagination.js')
        response.headers['Cache-Control'] = 'no-store, no-cache, must-revalidate, max-age=0'
        response.headers['Pragma'] = 'no-cache'
        response.headers['Expires'] = '0'
        return response
    except Exception as e:
        print(f"提供修复脚本时出错: {str(e)}")
        # 如果文件找不到，提供一个内联的基本修复脚本
        js_content = """
        // 基本分页修复脚本
        document.addEventListener('DOMContentLoaded', function() {
            console.log("运行应急修复脚本");
            // 防止空元素错误
            ['paginationControls', 'prevPage', 'nextPage', 'pageInfo'].forEach(id => {
                if (!document.getElementById(id)) {
                    console.log('创建缺失元素:' + id);
                    const el = document.createElement('div');
                    el.id = id;
                    el.style.display = 'none';
                    document.body.appendChild(el);
                }
            });
        });
        """
        response = app.response_class(response=js_content, status=200, mimetype='application/javascript')
        return response

# 在应用启动时初始化数据库
with app.app_context():
    check_db_tables()

if __name__ == '__main__':
    try:
        print("正在尝试启动在端口 5001...")
        app.debug = True  # 启用调试模式以获取更详细的错误信息
        app.run(host='0.0.0.0', port=5002)
    except OSError:
        print("端口 5001 被占用，尝试下一个...")
        app.debug = True  # 启用调试模式以获取更详细的错误信息
        app.run(host='0.0.0.0', port=5003)
