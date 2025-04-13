#!/usr/bin/env python
# -*- coding: utf-8 -*-

import os
import io
import base64
import sqlite3
import logging
import traceback
from PIL import Image, ImageDraw, ImageFont
import fitz  # PyMuPDF
import docx
from zipfile import ZipFile

class PaperPreviewGenerator:
    """试卷预览生成器 - 生成试卷预览缩略图和元数据"""
    
    def __init__(self, db_path, upload_folder):
        self.db_path = db_path
        self.upload_folder = upload_folder
        self.logger = logging.getLogger('paper_preview')
        self.logger.setLevel(logging.INFO)
        
        # 设置处理器
        if not self.logger.handlers:
            handler = logging.StreamHandler()
            formatter = logging.Formatter('%(asctime)s - %(name)s - %(levelname)s - %(message)s')
            handler.setFormatter(formatter)
            self.logger.addHandler(handler)
        
        # 支持的文件类型
        self.file_processors = {
            '.pdf': self.process_pdf,
            '.docx': self.process_docx,
            '.doc': self.process_doc,
            '.zip': self.process_zip,
        }
    
    def get_paper_data(self, paper_ids):
        """从数据库获取试卷信息"""
        results = {}
        
        try:
            conn = sqlite3.connect(self.db_path)
            cursor = conn.cursor()
            
            # 获取指定ID的试卷信息
            placeholders = ','.join(['?' for _ in paper_ids])
            query = f"""
                SELECT id, name, file_path
                FROM papers
                WHERE id IN ({placeholders})
            """
            
            cursor.execute(query, paper_ids)
            papers = cursor.fetchall()
            
            for paper_id, name, file_path in papers:
                try:
                    # 构建完整路径
                    full_path = os.path.join(self.upload_folder, file_path) if file_path else None
                    
                    # 检查文件是否存在
                    if full_path and os.path.exists(full_path):
                        # 获取文件扩展名
                        _, ext = os.path.splitext(full_path.lower())
                        
                        # 选择适当的处理器
                        processor = self.file_processors.get(ext)
                        
                        if processor:
                            # 处理文件
                            thumbnail, meta = processor(full_path)
                            
                            # 保存结果
                            results[str(paper_id)] = {
                                'thumbnail': thumbnail,
                                'meta': meta
                            }
                        else:
                            # 为不支持的文件类型创建默认预览
                            thumbnail, meta = self.create_default_preview(name, ext)
                            results[str(paper_id)] = {
                                'thumbnail': thumbnail,
                                'meta': meta
                            }
                    else:
                        # 文件不存在，创建默认预览
                        thumbnail, meta = self.create_default_preview(name, "缺失")
                        results[str(paper_id)] = {
                            'thumbnail': thumbnail,
                            'meta': meta
                        }
                except Exception as e:
                    self.logger.error(f"处理试卷ID={paper_id}失败: {str(e)}")
                    traceback.print_exc()
                    # 创建错误预览
                    thumbnail, meta = self.create_default_preview(name, "错误")
                    results[str(paper_id)] = {
                        'thumbnail': thumbnail,
                        'meta': meta
                    }
            
            conn.close()
        except Exception as e:
            self.logger.error(f"获取试卷数据失败: {str(e)}")
            traceback.print_exc()
        
        return results
    
    def process_pdf(self, file_path):
        """处理PDF文件"""
        try:
            # 打开PDF
            doc = fitz.open(file_path)
            
            # 获取页数
            page_count = len(doc)
            
            # 估计题目数量
            question_count = self.estimate_question_count_pdf(doc)
            
            # 生成第一页缩略图
            if page_count > 0:
                page = doc[0]
                pix = page.get_pixmap(matrix=fitz.Matrix(0.2, 0.2))
                
                # 转换为PIL图像
                img = Image.open(io.BytesIO(pix.tobytes("png")))
                
                # 处理缩略图
                thumb = self.process_thumbnail(img)
            else:
                # 创建默认缩略图
                thumb = self.create_text_thumbnail("PDF", "无页面")
            
            # 关闭文档
            doc.close()
            
            return thumb, {'pageCount': page_count, 'questionCount': question_count}
        except Exception as e:
            self.logger.error(f"处理PDF失败: {str(e)}")
            traceback.print_exc()
            return self.create_text_thumbnail("PDF", "错误"), {'error': str(e)}
    
    def process_docx(self, file_path):
        """处理DOCX文件"""
        try:
            # 打开DOCX
            doc = docx.Document(file_path)
            
            # 获取段落数
            para_count = len(doc.paragraphs)
            
            # 估计题目数量
            question_count = self.estimate_question_count_docx(doc)
            
            # 创建缩略图
            thumb = self.create_text_thumbnail("DOCX", f"{para_count}段")
            
            return thumb, {'paraCount': para_count, 'questionCount': question_count}
        except Exception as e:
            self.logger.error(f"处理DOCX失败: {str(e)}")
            traceback.print_exc()
            return self.create_text_thumbnail("DOCX", "错误"), {'error': str(e)}
    
    def process_doc(self, file_path):
        """处理DOC文件"""
        # 由于Python难以直接处理DOC文件，我们创建一个占位符缩略图
        return self.create_text_thumbnail("DOC", ""), {'format': 'DOC'}
    
    def process_zip(self, file_path):
        """处理ZIP文件"""
        try:
            # 打开ZIP
            with ZipFile(file_path, 'r') as zip_ref:
                # 获取文件列表
                file_list = zip_ref.namelist()
                file_count = len(file_list)
                
                # 创建缩略图
                thumb = self.create_text_thumbnail("ZIP", f"{file_count}文件")
                
                return thumb, {'fileCount': file_count, 'files': file_list[:5]}  # 只返回前5个文件名
        except Exception as e:
            self.logger.error(f"处理ZIP失败: {str(e)}")
            traceback.print_exc()
            return self.create_text_thumbnail("ZIP", "错误"), {'error': str(e)}
    
    def create_default_preview(self, name, type_info):
        """为不支持的文件类型创建默认预览"""
        return self.create_text_thumbnail(type_info, ""), {'name': name}
    
    def create_text_thumbnail(self, file_type, info_text):
        """创建文本缩略图"""
        # 创建一个白色背景的图像
        width, height = 100, 100
        img = Image.new('RGB', (width, height), color=(255, 255, 255))
        draw = ImageDraw.Draw(img)
        
        # 尝试加载字体，如果失败则使用默认字体
        try:
            font = ImageFont.truetype("Arial", 16)
            small_font = ImageFont.truetype("Arial", 12)
        except IOError:
            font = ImageFont.load_default()
            small_font = ImageFont.load_default()
        
        # 绘制文件类型
        text_width = draw.textlength(file_type, font=font)
        draw.text(((width - text_width) / 2, 30), file_type, font=font, fill=(0, 0, 0))
        
        # 绘制信息文本
        if info_text:
            info_width = draw.textlength(info_text, font=small_font)
            draw.text(((width - info_width) / 2, 60), info_text, font=small_font, fill=(100, 100, 100))
        
        # 绘制边框
        draw.rectangle([(0, 0), (width - 1, height - 1)], outline=(200, 200, 200))
        
        # 转换为base64
        buffer = io.BytesIO()
        img.save(buffer, format='PNG')
        img_str = base64.b64encode(buffer.getvalue()).decode('utf-8')
        
        return f"data:image/png;base64,{img_str}"
    
    def process_thumbnail(self, img):
        """处理缩略图"""
        # 确保缩略图大小一致
        img.thumbnail((100, 100), Image.LANCZOS)
        
        # 转换为base64
        buffer = io.BytesIO()
        img.save(buffer, format='PNG')
        img_str = base64.b64encode(buffer.getvalue()).decode('utf-8')
        
        return f"data:image/png;base64,{img_str}"
    
    def estimate_question_count_pdf(self, doc):
        """估计PDF中的题目数量"""
        question_count = 0
        
        try:
            # 检查页数限制以避免处理大型文件
            max_pages = min(5, len(doc))
            
            # 简单的题号检测模式
            patterns = [
                r'\d+\.', r'\d+、', r'\d+．',  # 中文题号样式
                r'第\s*\d+\s*题', r'题\s*\d+',  # 带"题"字的样式
                r'\(\s*\d+\s*\)', r'（\s*\d+\s*）'  # 带括号的题号
            ]
            
            # 只处理前几页以提高性能
            for page_num in range(max_pages):
                page = doc[page_num]
                text = page.get_text()
                
                # 计算每种模式的匹配数
                for pattern in patterns:
                    matches = len(page.search_for(pattern))
                    question_count += matches
            
            # 由于可能有重复计数，我们取一个合理的估计值
            question_count = question_count // (len(patterns) // 2)
        except Exception as e:
            self.logger.error(f"估计PDF题目数量失败: {str(e)}")
        
        return max(1, question_count)  # 至少返回1
    
    def estimate_question_count_docx(self, doc):
        """估计DOCX中的题目数量"""
        question_count = 0
        
        try:
            # 题号匹配的正则表达式模式
            import re
            
            patterns = [
                r'\d+\.', r'\d+、', r'\d+．',  # 中文题号样式
                r'第\s*\d+\s*题', r'题\s*\d+',  # 带"题"字的样式
                r'\(\s*\d+\s*\)', r'（\s*\d+\s*）'  # 带括号的题号
            ]
            
            # 合并所有段落文本
            full_text = ' '.join([p.text for p in doc.paragraphs])
            
            # 计算每种模式的匹配数
            for pattern in patterns:
                matches = len(re.findall(pattern, full_text))
                question_count += matches
            
            # 估计值
            question_count = question_count // (len(patterns) // 2)
        except Exception as e:
            self.logger.error(f"估计DOCX题目数量失败: {str(e)}")
        
        return max(1, question_count)  # 至少返回1
